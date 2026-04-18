#!/usr/bin/env python3
"""Solar Pro 3 API로 게임 기획 요소를 전문가 프레임워크로 분석합니다."""
import sys, os, json, argparse, re
import requests

try:
    from openai import OpenAI
except ImportError:
    print("Error: openai 패키지가 필요합니다. pip install openai", file=sys.stderr)
    sys.exit(1)

if "UPSTAGE_API_KEY" not in os.environ:
    print("Error: UPSTAGE_API_KEY 환경변수를 설정하세요", file=sys.stderr)
    print("  https://console.upstage.ai 에서 발급받을 수 있습니다", file=sys.stderr)
    sys.exit(1)

client = OpenAI(
    api_key=os.environ["UPSTAGE_API_KEY"],
    base_url="https://api.upstage.ai/v1"
)

# ── 전문가별 시스템 프롬프트 ──

EXPERT_PROMPTS = {
    "schell": {
        "name": "Jesse Schell — 구조 분석",
        "system": """IMPORTANT: 반드시 한국어로만 답변하세요. 영어를 절대 사용하지 마세요.

당신은 Jesse Schell의 "The Art of Game Design" 방법론을 따르는 게임 분석가입니다.
게임을 Elemental Tetrad(Mechanics, Story, Aesthetics, Technology) 4축으로 분해하고,
각 축이 서로를 어떻게 지지하는지 평가하세요.
이어서 다음 렌즈를 적용하세요:
1. Essential Experience: 본질적 경험은?
2. Emergence: 단순한 규칙이 만드는 복잡한 행동은?
3. Interest Curve: 경험 강도의 시간적 흐름은?
4. Flow: 몰입 조건의 충족 여부는?
5. Theme: 테마와 메커닉의 공명 수준은?
한국어로 분석하세요."""
    },
    "koster": {
        "name": "Raph Koster — 재미 구조",
        "system": """IMPORTANT: 반드시 한국어로만 답변하세요. 영어를 절대 사용하지 마세요.

당신은 Raph Koster의 "A Theory of Fun" 방법론을 따르는 게임 분석가입니다.
Koster의 핵심 명제: "재미란 패턴을 학습하는 과정에서 오는 쾌감"
다음 관점에서 분석하세요:
1. 핵심 패턴 식별: 플레이어가 학습하는 주요 패턴 3~5개
2. 패턴 깊이: 각 패턴의 복잡도와 숙달 소요 시간
3. 학습 곡선: 패턴이 소개되고 심화되는 순서
4. 소진 분석: 패턴 grok 후 흥미 유지 메커니즘
5. Chunking: 복잡한 시스템을 학습 가능한 단위로 분해하는 방식
한국어로 분석하세요."""
    },
    "mda": {
        "name": "MDA 프레임워크 — 시스템 인과",
        "system": """IMPORTANT: 반드시 한국어로만 답변하세요. 영어를 절대 사용하지 마세요.

당신은 Hunicke, LeBlanc, Zubek의 MDA 프레임워크를 따르는 게임 분석가입니다.
게임을 MDA 3계층으로 분해하세요:
1. Aesthetics 식별: 8가지 미학 유형(Sensation, Fantasy, Narrative, Challenge,
   Fellowship, Discovery, Expression, Submission) 중 주력 2~3개
2. Dynamics 매핑: 각 주력 Aesthetic을 지지하는 플레이어 행동 패턴
3. Mechanics 추적: 각 Dynamics를 가능하게 하는 구체적 규칙/시스템
4. 인과 체인 평가: M→D→A 각 전환의 자연스러움
5. 관점 비대칭: 디자이너 의도 vs 플레이어 실제 경험의 간극
한국어로 분석하세요."""
    },
    "sylvester": {
        "name": "Tynan Sylvester — 감정 공학",
        "system": """IMPORTANT: 반드시 한국어로만 답변하세요. 영어를 절대 사용하지 마세요.

당신은 Tynan Sylvester의 "Designing Games" 방법론을 따르는 게임 분석가입니다.
핵심 명제: "게임은 경험을 생성하는 기계이며, 경험은 감정의 흐름이다."
다음 관점에서 분석하세요:
1. 감정 체인: Mechanics → Fiction → Events → Emotions → Experience
2. 감정 트리거: 10가지(Learning, Character Arc, Challenge, Social, Acquisition,
   Music, Spectacle, Beauty, Environment, Primal Threats) 중 활용되는 것
3. 가치 변화: 안전↔위험, 고립↔유대 등 인간 가치 변화의 폭
4. 미귀인 효과: 메커닉의 각성과 픽션의 라벨링 조화
5. 엘레강스: 메커닉 수 대비 의미 있는 전략/경험 비율
한국어로 분석하세요."""
    },
    "miyamoto": {
        "name": "Shigeru Miyamoto — 접근성 설계",
        "system": """IMPORTANT: 반드시 한국어로만 답변하세요. 영어를 절대 사용하지 마세요.

당신은 Miyamoto의 직관적 게임 디자인 철학을 따르는 게임 분석가입니다.
핵심 원칙: "환경이 가르쳐야 한다" + "게임은 보편적으로 접근 가능해야 한다"
다음 관점에서 분석하세요:
1. 첫인상: 게임의 첫 5분이 어떤 경험을 제공하는가
2. 반응성: 핵심 조작의 즉각성과 직관성
3. A-ha! 모먼트: 스스로 발견하도록 설계된 순간
4. 온보딩 곡선: 초보자 진입 과정의 매끄러움, 이탈 지점
5. 기술-경험 통합: 플랫폼 특성과 게임 경험의 관계
한국어로 분석하세요."""
    },
    "miyazaki": {
        "name": "Hidetaka Miyazaki — 극복의 미학",
        "system": """IMPORTANT: 반드시 한국어로만 답변하세요. 영어를 절대 사용하지 마세요.

당신은 Miyazaki의 환경 스토리텔링과 극복 미학을 따르는 게임 분석가입니다.
핵심 명제: "세계 자체가 이야기를 말해야 한다" + "높은 난이도는 보상의 전제 조건"
다음 관점에서 분석하세요:
1. 환경 서사: 세계가 스토리를 전달하는 방식 (명시적 vs 환경적 비율)
2. 난이도 곡선: 도전→학습→극복→보상 사이클의 건강함
3. 발견 설계: 탐색과 발견의 보상 구조
4. 모호함의 가치: 의도적으로 설명하지 않는 것들의 효과
5. 감정 아크: 좌절과 극복이 만드는 전체적 감정 곡선
한국어로 분석하세요."""
    }
}

INTEGRATED_SYSTEM = """IMPORTANT: 반드시 한국어로만 답변하세요. 영어를 절대 사용하지 마세요.

당신은 세계적 게임 디자인 이론에 정통한 전문 게임 분석가입니다.
아래 6개 전문가의 개별 분석 결과를 읽고 교차 분석을 수행하세요.
- 다수 전문가가 동일하게 지적하는 강점 3개
- 다수 전문가가 동일하게 지적하는 약점 3개
- 전문가 간 의견이 갈리는 지점과 그 이유
- 기획 관점에서의 핵심 혁신 포인트
한국어로 작성하세요."""

GAP_ANALYSIS_SYSTEM = """IMPORTANT: 반드시 한국어로만 답변하세요. 영어를 절대 사용하지 마세요.

당신은 게임 디자인 이론과 플레이어 리뷰 데이터를 교차 분석하는 전문가입니다.

아래에 두 가지 자료가 주어집니다:
1. 전문가 프레임워크 분석 결과 (이론적 관점)
2. Steam 리뷰 패턴 분석 결과 (실제 플레이어 반응)

이 둘을 대조하여 **디자이너 의도와 플레이어 경험 사이의 간극**을 분석하세요.

## 분석 구조

### 미학 간극
MDA 분석이 예측한 주력 Aesthetic(8가지 중)과 리뷰에서 실제로 칭찬/불만이 집중되는 요소를 대조.
일치하면 "디자인 의도가 잘 전달됨", 불일치하면 그 원인을 추론.

### 패턴 간극
Koster 분석이 식별한 핵심 학습 패턴과 리뷰에서 플레이어가 실제로 "재밌다"고 언급하는 요소를 대조.
이론적 핵심 패턴이 실제로 재미의 원천인지, 아니면 다른 요소가 재미를 주고 있는지.

### 접근성 간극
Miyamoto 분석이 평가한 온보딩/접근성 수준과 부정 리뷰에서 나타나는 진입장벽 불만을 대조.
이론상 문제없어 보이지만 플레이어가 실제로 겪는 허들은 무엇인지.

### 감정 간극
Sylvester 분석이 예측한 감정 트리거와 리뷰에서 플레이어가 보고하는 실제 감정을 대조.
의도한 감정(예: 도전의 만족)과 실제 감정(예: 좌절)의 차이.

### 간극에서 도출되는 기획 시사점
위 간극들에서 도출되는 핵심 기획 시사점 3개를 정리.
각 시사점에 대해 "개선 방향"을 1줄로 제안.

한국어로 작성하세요."""

STEAM_HEADERS = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}


def resolve_steam_appid(query: str) -> str:
    """Steam URL, 숫자 ID, 또는 게임 이름을 받아 app ID를 반환한다."""
    # URL에서 추출
    m = re.search(r'/app/(\d+)', query)
    if m:
        return m.group(1)
    # 숫자 ID 그대로
    if query.strip().isdigit():
        return query.strip()
    # 게임 이름으로 검색
    print(f"[Steam] '{query}' 검색 중...", file=sys.stderr)
    r = requests.get(
        "https://store.steampowered.com/api/storesearch",
        params={"term": query, "cc": "US", "l": "english"},
        headers=STEAM_HEADERS, timeout=15
    )
    items = r.json().get("items", [])
    if not items:
        raise ValueError(f"Steam에서 '{query}'를 찾을 수 없습니다.")
    appid = str(items[0]["id"])
    print(f"[Steam] '{items[0]['name']}' (ID: {appid}) 발견", file=sys.stderr)
    return appid


def fetch_steam_game_data(query: str) -> tuple[str, str | None]:
    """Steam URL / ID / 게임명으로부터 게임 정보와 리뷰 요약을 가져온다."""
    appid = resolve_steam_appid(query)
    print(f"[Steam] App ID: {appid} 상세 정보 수집 중...", file=sys.stderr)

    # ── 게임 상세 정보 ──
    resp = requests.get(
        "https://store.steampowered.com/api/appdetails",
        params={"appids": appid, "cc": "US"},
        headers=STEAM_HEADERS, timeout=15
    )
    resp.raise_for_status()
    detail_json = resp.json()

    if not detail_json.get(appid, {}).get("success"):
        raise RuntimeError(f"Steam API에서 app {appid} 정보를 가져오지 못했습니다.")

    data        = detail_json[appid]["data"]
    genres      = ", ".join(g["description"] for g in data.get("genres", []))
    categories  = ", ".join(c["description"] for c in data.get("categories", []))
    platforms   = [k for k, v in data.get("platforms", {}).items() if v]
    developers  = ", ".join(data.get("developers", []))
    publishers  = ", ".join(data.get("publishers", []))
    release     = data.get("release_date", {}).get("date", "")
    description = data.get("detailed_description", "") or data.get("short_description", "")
    description = re.sub(r'<[^>]+>', ' ', description).strip()
    description = re.sub(r'\s+', ' ', description)
    if len(description) > 2000:
        description = description[:2000] + "..."
    metacritic      = data.get("metacritic", {}).get("score", "없음")
    recommendations = data.get("recommendations", {}).get("total", "없음")

    game_data = json.dumps({
        "title":           data.get("name", ""),
        "developer":       developers,
        "publisher":       publishers,
        "genre":           genres,
        "categories":      categories,
        "platform":        ", ".join(platforms),
        "release":         release,
        "metacritic_score": metacritic,
        "total_reviews":   recommendations,
        "description":     description,
        "steam_url":       f"https://store.steampowered.com/app/{appid}/",
    }, ensure_ascii=False, indent=2)

    print(f"[Steam] '{data.get('name')}' 정보 수집 완료", file=sys.stderr)

    # ── 리뷰 수집 ──
    print(f"[Steam] 한국어 리뷰 수집 중...", file=sys.stderr)
    r2 = requests.get(
        f"https://store.steampowered.com/appreviews/{appid}",
        params={"json": 1, "language": "koreana", "num_per_page": 40,
                "review_type": "all", "purchase_type": "steam", "filter": "recent"},
        headers=STEAM_HEADERS, timeout=15
    )
    review_summary = None
    if r2.ok:
        reviews  = r2.json().get("reviews", [])
        pos = [rv["review"] for rv in reviews if rv.get("voted_up")]
        neg = [rv["review"] for rv in reviews if not rv.get("voted_up")]

        def trim(lst, n=6):
            return [re.sub(r'\s+', ' ', r).strip()[:300] for r in lst[:n] if r.strip()]

        parts = []
        if pos: parts.append("【긍정 리뷰】\n" + "\n".join(f"- {r}" for r in trim(pos)))
        if neg: parts.append("【부정 리뷰】\n" + "\n".join(f"- {r}" for r in trim(neg)))
        if parts:
            review_summary = "\n\n".join(parts)
            print(f"[Steam] 리뷰 수집 완료 (긍정 {len(pos)}, 부정 {len(neg)})", file=sys.stderr)

    return game_data, review_summary


def strip_english_cot(text: str) -> str:
    """Solar Pro 3가 앞에 출력하는 영어 chain-of-thought를 제거하고 한국어 본문만 반환한다."""
    lines = text.split("\n")
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        # 한글 문자가 포함된 줄이 나오면 그 지점부터 반환
        if re.search(r'[\uAC00-\uD7A3]', stripped):
            return "\n".join(lines[i:])
    return text  # 한글을 못 찾으면 원본 반환


EXPERT_ORDER = ["schell", "koster", "mda", "sylvester", "miyamoto", "miyazaki"]
EXPERT_LABELS = {
    "schell": "Jesse Schell — 구조 분석",
    "koster": "Raph Koster — 재미 구조",
    "mda": "MDA 프레임워크 — 시스템 인과",
    "sylvester": "Tynan Sylvester — 감정 공학",
    "miyamoto": "Shigeru Miyamoto — 접근성 설계",
    "miyazaki": "Hidetaka Miyazaki — 극복의 미학",
    "integrated": "통합 교차 분석",
    "gap": "디자이너 의도 vs 플레이어 경험 간극 분석",
}


def analyze(game_data: str, expert: str, review_summary: str = None) -> str:
    """단일 전문가 관점으로 분석을 수행한다. 리뷰 요약이 있으면 보조 자료로 주입."""
    prompt_info = EXPERT_PROMPTS[expert]

    user_content = f"다음 게임을 분석하세요:\n\n{game_data}"
    if review_summary:
        user_content += (
            f"\n\n---\n"
            f"[참고: 실제 Steam 리뷰 패턴 요약]\n"
            f"아래는 실제 플레이어 리뷰를 분석한 결과입니다. "
            f"이론적 분석과 실제 플레이어 반응을 대조하면서 분석하세요.\n\n"
            f"{review_summary}"
        )

    resp = client.chat.completions.create(
        model="solar-pro3",
        messages=[
            {"role": "system", "content": prompt_info["system"]},
            {"role": "user", "content": user_content},
            {"role": "assistant", "content": "## 분석 결과\n\n"},
        ],
        temperature=0.7,
        max_tokens=4096
    )
    content = strip_english_cot(resp.choices[0].message.content)
    return "## 분석 결과\n\n" + content


def integrate(game_data: str, individual_results: dict) -> str:
    """6개 분석 결과를 통합하여 교차 분석을 생성한다."""
    combined = f"## 분석 대상\n{game_data}\n\n"
    for expert, result in individual_results.items():
        if expert in EXPERT_PROMPTS:
            name = EXPERT_PROMPTS[expert]["name"]
            combined += f"## {name}\n{result}\n\n"

    resp = client.chat.completions.create(
        model="solar-pro3",
        messages=[
            {"role": "system", "content": INTEGRATED_SYSTEM},
            {"role": "user", "content": combined},
            {"role": "assistant", "content": "## 통합 교차 분석\n\n"},
        ],
        temperature=0.5,
        max_tokens=4096
    )
    content = strip_english_cot(resp.choices[0].message.content)
    return "## 통합 교차 분석\n\n" + content


def analyze_gap(individual_results: dict, review_summary: str) -> str:
    """전문가 분석과 리뷰 데이터를 대조하여 의도-경험 간극을 분석한다."""
    combined = "## 전문가 프레임워크 분석 결과\n\n"
    for expert, result in individual_results.items():
        if expert in EXPERT_PROMPTS:
            name = EXPERT_PROMPTS[expert]["name"]
            combined += f"### {name}\n{result}\n\n"

    combined += f"---\n\n## Steam 리뷰 패턴 분석 결과\n\n{review_summary}"

    resp = client.chat.completions.create(
        model="solar-pro3",
        messages=[
            {"role": "system", "content": GAP_ANALYSIS_SYSTEM},
            {"role": "user", "content": combined},
            {"role": "assistant", "content": "## 간극 분석\n\n"},
        ],
        temperature=0.5,
        max_tokens=4096
    )
    content = strip_english_cot(resp.choices[0].message.content)
    return "## 간극 분석\n\n" + content


def results_to_markdown(game_title: str, results: dict) -> str:
    """분석 결과 딕셔너리를 마크다운 문서로 변환한다."""
    md = f"# 게임 분석 리포트: {game_title}\n\n"
    md += "---\n\n"

    order = EXPERT_ORDER + ["integrated", "gap"]
    for key in order:
        if key not in results:
            continue
        label = EXPERT_LABELS.get(key, key)
        md += f"# {label}\n\n"
        md += results[key]
        md += "\n\n---\n\n"

    return md


def save_output(results: dict, output_path: str, game_title: str = "게임"):
    """확장자에 따라 PDF / DOCX / MD / JSON 으로 저장한다."""
    ext = os.path.splitext(output_path)[1].lower()

    if ext == ".pdf":
        try:
            from fpdf import FPDF
        except ImportError:
            print("PDF 저장에는 fpdf2가 필요합니다: pip install fpdf2", file=sys.stderr)
            sys.exit(1)

        md_text = results_to_markdown(game_title, results)

        pdf = FPDF()
        pdf.add_page()

        # 한글 폰트 경로 탐색 (Windows 기본 폰트)
        font_candidates = [
            "C:/Windows/Fonts/malgun.ttf",       # 맑은 고딕
            "C:/Windows/Fonts/NanumGothic.ttf",  # 나눔고딕 (설치된 경우)
        ]
        font_path = next((p for p in font_candidates if os.path.exists(p)), None)
        if font_path is None:
            print("Error: 한글 폰트를 찾을 수 없습니다. malgun.ttf 또는 NanumGothic.ttf가 필요합니다.", file=sys.stderr)
            sys.exit(1)

        pdf.add_font("Korean", "", font_path, uni=True)
        pdf.set_font("Korean", size=11)
        pdf.set_auto_page_break(auto=True, margin=15)

        for line in md_text.split("\n"):
            # 마크다운 헤딩 처리
            if line.startswith("# "):
                pdf.set_font("Korean", size=16)
                pdf.multi_cell(0, 10, line[2:])
                pdf.set_font("Korean", size=11)
            elif line.startswith("## "):
                pdf.set_font("Korean", size=14)
                pdf.multi_cell(0, 9, line[3:])
                pdf.set_font("Korean", size=11)
            elif line.startswith("### "):
                pdf.set_font("Korean", size=12)
                pdf.multi_cell(0, 8, line[4:])
                pdf.set_font("Korean", size=11)
            elif line.strip() == "---":
                pdf.ln(4)
                pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
                pdf.ln(4)
            else:
                pdf.multi_cell(0, 7, line)

        pdf.output(output_path)

    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
        except ImportError:
            print("DOCX 저장에는 python-docx가 필요합니다: pip install python-docx", file=sys.stderr)
            sys.exit(1)

        md_text = results_to_markdown(game_title, results)
        doc = Document()
        doc.core_properties.title = f"게임 분석 리포트: {game_title}"

        for line in md_text.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.strip() == "---":
                doc.add_paragraph("─" * 50)
            elif line.strip():
                doc.add_paragraph(line)

        doc.save(output_path)

    elif ext == ".md":
        md_text = results_to_markdown(game_title, results)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_text)

    else:
        # 기본: JSON
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(results, f, ensure_ascii=False, indent=2)


def main():
    parser = argparse.ArgumentParser(description="게임 기획 분석기")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--game-data", help="게임 정보 JSON 파일 경로")
    group.add_argument("--steam", help="Steam 링크 또는 App ID (예: https://store.steampowered.com/app/1177620/)")
    parser.add_argument("--review-data", default=None, help="리뷰 JSON 파일 경로 (--game-data 사용 시, 선택)")
    parser.add_argument("--analysis", default="all",
                        help="분석 유형: all, gap, schell, koster, mda, sylvester, miyamoto, miyazaki")
    parser.add_argument("--output", default=None,
                        help="결과 저장 경로 (.pdf / .docx / .md / .json)")
    args = parser.parse_args()

    review_summary = None

    if args.steam:
        # Steam URL/ID에서 직접 수집
        try:
            game_data, review_summary = fetch_steam_game_data(args.steam)
        except Exception as e:
            print(f"Error: Steam 수집 실패 — {e}", file=sys.stderr)
            sys.exit(1)
    else:
        # 기존 JSON 파일 방식
        with open(args.game_data, "r", encoding="utf-8") as f:
            game_data = f.read()

        if args.review_data and os.path.isfile(args.review_data):
            with open(args.review_data, "r", encoding="utf-8") as f:
                review_json = json.load(f)
                review_summary = review_json.get("analysis", "")
            print(f"[리뷰 데이터 로드] {args.review_data}", file=sys.stderr)

    # 게임 제목 추출
    try:
        game_title = json.loads(game_data).get("title", "Unknown")
    except Exception:
        game_title = os.path.splitext(os.path.basename(args.game_data or "game"))[0]

    # 간극 분석만 요청된 경우
    if args.analysis == "gap":
        if not review_summary:
            print("Error: 간극 분석에는 --review-data가 필요합니다", file=sys.stderr)
            sys.exit(1)
        experts = list(EXPERT_PROMPTS.keys())
        results = {}
        for expert in experts:
            name = EXPERT_PROMPTS[expert]["name"]
            print(f"[분석 중] {name}...", file=sys.stderr)
            results[expert] = analyze(game_data, expert, review_summary)
            print(f"[완료] {name}", file=sys.stderr)

        print("[간극 분석 중]...", file=sys.stderr)
        results["gap"] = analyze_gap(results, review_summary)
        print("[간극 분석 완료]", file=sys.stderr)

        if args.output:
            save_output(results, args.output, game_title)
            print(f"결과 저장: {args.output}", file=sys.stderr)
        else:
            print(results_to_markdown(game_title, results))
        return

    # 일반 분석
    if args.analysis == "all":
        experts = list(EXPERT_PROMPTS.keys())
    else:
        experts = [args.analysis]

    results = {}
    for expert in experts:
        name = EXPERT_PROMPTS[expert]["name"]
        print(f"[분석 중] {name}...", file=sys.stderr)
        results[expert] = analyze(game_data, expert, review_summary)
        print(f"[완료] {name}", file=sys.stderr)

    if args.analysis == "all":
        print("[통합 분석 중]...", file=sys.stderr)
        results["integrated"] = integrate(game_data, results)
        print("[통합 분석 완료]", file=sys.stderr)

        if review_summary:
            print("[간극 분석 중]...", file=sys.stderr)
            results["gap"] = analyze_gap(results, review_summary)
            print("[간극 분석 완료]", file=sys.stderr)

    if args.output:
        save_output(results, args.output, game_title)
        print(f"결과 저장: {args.output}", file=sys.stderr)
    else:
        print(results_to_markdown(game_title, results))


if __name__ == "__main__":
    main()
